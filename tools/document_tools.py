"""
Document Tools

Used by: WriterAgent, API routes
Responsibilities:
- Convert LLM-generated text into professional PDF/DOCX documents
- Extract text from uploaded documents (PDF, DOCX, TXT)
- Format with headers, sections, citations
- Save to output directory
"""
import re
from datetime import datetime
from pathlib import Path


def generate_pdf(
    content: str,
    title: str,
    output_path: str | Path,
    sector: str = "",
    metadata: dict | None = None,
) -> Path:
    """
    Generate a professional PDF from text content using ReportLab.

    Args:
        content: Document body text (markdown-like: # headers, ** bold)
        title: Document title shown in header
        output_path: Where to save the PDF
        sector: Sector name for the footer
        metadata: Optional dict with "author", "company", "date" keys

    Returns:
        Path to the generated PDF
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    )
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    metadata = metadata or {}
    date_str = metadata.get("date", datetime.now().strftime("%d %B %Y"))
    author = metadata.get("author", "GRaC Compliance System")
    company = metadata.get("company", "")

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        rightMargin=2.5 * cm,
        leftMargin=2.5 * cm,
        topMargin=3 * cm,
        bottomMargin=2.5 * cm,
        title=title,
        author=author,
    )

    base_styles = getSampleStyleSheet()

    styles = {
        "title": ParagraphStyle(
            "DocTitle",
            parent=base_styles["Heading1"],
            fontSize=18,
            spaceAfter=6,
            textColor=colors.HexColor("#1a1a2e"),
            alignment=TA_CENTER,
        ),
        "subtitle": ParagraphStyle(
            "DocSubtitle",
            parent=base_styles["Normal"],
            fontSize=10,
            spaceAfter=20,
            textColor=colors.grey,
            alignment=TA_CENTER,
        ),
        "h1": ParagraphStyle(
            "H1",
            parent=base_styles["Heading2"],
            fontSize=13,
            spaceBefore=14,
            spaceAfter=6,
            textColor=colors.HexColor("#1a1a2e"),
            borderPad=4,
        ),
        "h2": ParagraphStyle(
            "H2",
            parent=base_styles["Heading3"],
            fontSize=11,
            spaceBefore=10,
            spaceAfter=4,
            textColor=colors.HexColor("#16213e"),
        ),
        "body": ParagraphStyle(
            "Body",
            parent=base_styles["Normal"],
            fontSize=10,
            leading=15,
            spaceAfter=8,
            alignment=TA_JUSTIFY,
        ),
        "bullet": ParagraphStyle(
            "Bullet",
            parent=base_styles["Normal"],
            fontSize=10,
            leading=14,
            leftIndent=18,
            spaceAfter=4,
        ),
    }

    story = []

    # Title block
    story.append(Paragraph(title, styles["title"]))
    subtitle_parts = [date_str]
    if company:
        subtitle_parts.insert(0, company)
    if sector:
        subtitle_parts.append(f"Sector: {sector.replace('_', ' ').title()}")
    story.append(Paragraph(" · ".join(subtitle_parts), styles["subtitle"]))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#1a1a2e")))
    story.append(Spacer(1, 0.4 * cm))

    # Parse content
    for line in content.split("\n"):
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 0.2 * cm))
            continue

        if stripped.startswith("## "):
            story.append(Paragraph(_escape(stripped[3:]), styles["h1"]))
        elif stripped.startswith("### "):
            story.append(Paragraph(_escape(stripped[4:]), styles["h2"]))
        elif stripped.startswith("# "):
            story.append(Paragraph(_escape(stripped[2:]), styles["h1"]))
        elif stripped.startswith("- ") or stripped.startswith("• "):
            bullet_text = _inline_markup(stripped[2:])
            story.append(Paragraph(f"• {bullet_text}", styles["bullet"]))
        elif re.match(r"^\d+\.\s", stripped):
            story.append(Paragraph(_inline_markup(stripped), styles["bullet"]))
        else:
            story.append(Paragraph(_inline_markup(stripped), styles["body"]))

    doc.build(story)
    return output_path


def generate_docx(
    content: str,
    title: str,
    output_path: str | Path,
    sector: str = "",
    metadata: dict | None = None,
) -> Path:
    """
    Generate a Word document from text content using python-docx.

    Args and Returns: same as generate_pdf()
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    metadata = metadata or {}
    date_str = metadata.get("date", datetime.now().strftime("%d %B %Y"))

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    sub = doc.add_paragraph(date_str)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(10)
    sub.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph()  # Spacer

    for line in content.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue

        if stripped.startswith("## ") or stripped.startswith("# "):
            text = stripped.lstrip("# ").strip()
            doc.add_heading(text, level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("- ") or stripped.startswith("• "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", stripped):
            doc.add_paragraph(stripped, style="List Number")
        else:
            p = doc.add_paragraph()
            _add_inline_runs(p, stripped)

    doc.save(str(output_path))
    return output_path


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _escape(text: str) -> str:
    """Escape ReportLab-unsafe XML characters."""
    return (
        text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
    )


def _inline_markup(text: str) -> str:
    """Convert **bold** and *italic* markdown to ReportLab XML tags."""
    text = _escape(text)
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"\*(.+?)\*", r"<i>\1</i>", text)
    return text


def _add_inline_runs(paragraph, text: str) -> None:
    """Add bold/italic inline markdown as separate runs to a docx paragraph."""
    import re as _re
    parts = _re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


# ---------------------------------------------------------------------------
# Document Reading (input)
# ---------------------------------------------------------------------------

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx"}


def extract_text_from_pdf(file_path: str | Path) -> str:
    """Extract text from a PDF file using pdfplumber (fallback PyMuPDF)."""
    file_path = Path(file_path)
    text = ""
    method = "none"
    import pdfplumber
    try:
        with pdfplumber.open(str(file_path)) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
            text = "\n".join(pages)
            method = "pdfplumber"
    except Exception:
        pass

    if not text.strip():
        try:
            import fitz
            doc = fitz.open(str(file_path))
            pages = [page.get_text() for page in doc]
            text = "\n".join(pages)
            method = "pymupdf"
            doc.close()
        except Exception:
            pass

    return text.strip()


def extract_text_from_docx(file_path: str | Path) -> str:
    """Extract text from a DOCX file using python-docx."""
    from docx import Document
    file_path = Path(file_path)
    doc = Document(str(file_path))
    paragraphs = [p.text for p in doc.paragraphs]
    return "\n".join(paragraphs).strip()


def extract_text_from_txt(file_path: str | Path) -> str:
    """Extract text from a plain text file."""
    file_path = Path(file_path)
    return file_path.read_text(encoding="utf-8", errors="replace").strip()


def read_document(file_path: str | Path) -> dict:
    """
    Read a document file and return its text content.

    Supported formats: PDF, DOCX, TXT

    Args:
        file_path: Path to the document file

    Returns:
        {"text": str, "format": str, "pages": int|None, "filename": str}

    Raises:
        ValueError: If the file format is not supported
    """
    file_path = Path(file_path)
    ext = file_path.suffix.lower()

    if ext == ".pdf":
        text = extract_text_from_pdf(file_path)
        fmt = "pdf"
    elif ext == ".docx":
        text = extract_text_from_docx(file_path)
        fmt = "docx"
    elif ext == ".txt":
        text = extract_text_from_txt(file_path)
        fmt = "txt"
    else:
        raise ValueError(f"Unsupported file format: {ext}. Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}")

    return {
        "text": text,
        "format": fmt,
        "filename": file_path.name,
    }
